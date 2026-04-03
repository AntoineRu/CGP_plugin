#!/usr/bin/env python3
"""
PostToolUse hook — output_router.py

Triggered on every Write/Edit. If the written file is a .md matching the
cgp-* convention, converts it to .docx and moves it to ~/CGP/<routing>/,
then deletes the source .md.

Stdin payload: {"tool_name": "Write", "tool_input": {"file_path": "..."}, ...}
Stdout: nothing (passive hook)
"""

import json
import pathlib
import re
import sys

# Files to route must have this prefix
CGP_PREFIX = "cgp-"
TRIGGER_EXTENSIONS = {".md"}

# Maps filename keyword → (subfolder_in_CGP, is_cabinet)
TYPE_MAP = [
    ("bilan",       ("bilans",          False)),
    ("rediger",     ("lettres",         False)),
    ("lettre",      ("lettres",         False)),
    ("mission",     ("lettres",         False)),
    ("-cr-",        ("lettres",         False)),
    ("analyser",    ("analyses",        False)),
    ("analyse",     ("analyses",        False)),
    ("rdv",         ("rendez-vous",     False)),
    ("rendez-vous", ("rendez-vous",     False)),
    ("reporting",   ("reporting",       False)),
    ("dossier",     ("reporting",       False)),
    ("rapport",     ("reporting",       False)),
    ("veille",      ("veille",          True)),
    ("vulgariser",  ("vulgarisation",   True)),
    ("vulgaris",    ("vulgarisation",   True)),
    ("marketing",   ("marketing",       True)),
    ("prospecter",  ("prospection",     True)),
    ("prospect",    ("prospection",     True)),
]

DEFAULT_CONFIG = {
    "template_path": None,
    "style_map": {
        "Heading 1": "Heading 1",
        "Heading 2": "Heading 2",
        "Heading 3": "Heading 3",
        "Normal": "Normal",
        "List Bullet": "List Bullet",
        "List Number": "List Number",
    },
    "typography": {
        "body_font": "Calibri",
        "body_size": 11,
        "heading_font": "Calibri",
        "heading_size": 14,
    },
}


def load_or_create_config() -> dict:
    config_path = pathlib.Path(__file__).parent / "charte_config.json"
    if config_path.exists():
        try:
            return json.loads(config_path.read_text(encoding="utf-8"))
        except Exception:
            pass
    # Create default config
    config_path.write_text(
        json.dumps(DEFAULT_CONFIG, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    sys.stderr.write("[cgp] charte_config.json absent — paramètres par défaut utilisés\n")
    return DEFAULT_CONFIG


def detect_type(stem: str) -> tuple[str, bool]:
    """Return (subfolder, is_cabinet). Defaults to ('divers', True)."""
    lower = stem.lower()
    for keyword, result in TYPE_MAP:
        if keyword in lower:
            return result
    return ("divers", True)


def detect_client(md_content: str) -> str | None:
    """Search md_content for known real names from the RGPD registry.
    Returns the real name (for folder naming) or None if not found."""
    registry_path = pathlib.Path.home() / ".cgp-client-registry.json"
    if not registry_path.exists():
        return None
    try:
        registry = json.loads(registry_path.read_text(encoding="utf-8"))
    except Exception:
        return None
    real_names = list(registry.get("real_to_pseudo", {}).keys())
    # Longest match first to avoid partial matches
    real_names.sort(key=len, reverse=True)
    for name in real_names:
        if name.lower() in md_content.lower():
            return name
    return None


def _apply_inline(para, text: str) -> None:
    """Apply **bold** and *italic* inline markers."""
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            para.add_run(part[2:-2]).bold = True
        elif part.startswith("*") and part.endswith("*"):
            para.add_run(part[1:-1]).italic = True
        else:
            para.add_run(part)


def _md_to_docx(md_path: pathlib.Path, out_path: pathlib.Path, config: dict) -> None:
    """Convert a Markdown file to .docx at out_path using config."""
    from docx import Document  # type: ignore

    template_path = config.get("template_path")
    style_map: dict = config.get("style_map", {})

    def style(key: str) -> str:
        return style_map.get(key, key)

    if template_path:
        tp = pathlib.Path(template_path)
        if not tp.exists():
            raise FileNotFoundError(f"Template introuvable : {tp}")
        doc = Document(str(tp))
        # Clear body content, keep styles/headers/footers
        body = doc.element.body
        for child in list(body):
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag in ("p", "tbl"):
                body.remove(child)
    else:
        doc = Document()

    lines = md_path.read_text(encoding="utf-8").splitlines()

    for line in lines:
        stripped = line.rstrip()

        if stripped.startswith("### "):
            doc.add_paragraph(stripped[4:], style=style("Heading 3"))
        elif stripped.startswith("## "):
            doc.add_paragraph(stripped[3:], style=style("Heading 2"))
        elif stripped.startswith("# "):
            doc.add_paragraph(stripped[2:], style=style("Heading 1"))
        elif re.match(r"^[-*] ", stripped):
            doc.add_paragraph(stripped[2:], style=style("List Bullet"))
        elif re.match(r"^\d+\. ", stripped):
            doc.add_paragraph(re.sub(r"^\d+\. ", "", stripped), style=style("List Number"))
        elif re.match(r"^---+$", stripped):
            doc.add_paragraph("", style=style("Normal"))
        elif stripped.startswith("|"):
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            if not all(re.match(r"^[-:]+$", c) for c in cells):
                doc.add_paragraph("\t".join(cells), style=style("Normal"))
        elif stripped == "":
            pass
        else:
            para = doc.add_paragraph(style=style("Normal"))
            _apply_inline(para, stripped)

    doc.save(str(out_path))


def main() -> None:
    try:
        payload = json.load(sys.stdin)
    except Exception:
        return

    if payload.get("tool_name") not in ("Write", "Edit"):
        return

    file_path = payload.get("tool_input", {}).get("file_path", "")
    md_path = pathlib.Path(file_path)

    if md_path.suffix not in TRIGGER_EXTENSIONS:
        return
    if not md_path.name.startswith(CGP_PREFIX):
        return
    if not md_path.exists():
        return

    config = load_or_create_config()
    md_content = md_path.read_text(encoding="utf-8")

    subfolder, is_cabinet = detect_type(md_path.stem)
    client_name = None if is_cabinet else detect_client(md_content)

    # Routing
    cgp_root = pathlib.Path.home() / "CGP"
    if is_cabinet or client_name is None:
        target_dir = cgp_root / "_cabinet" / subfolder
    else:
        # Sanitise client name for filesystem
        safe_name = re.sub(r"[^\w\s\-]", "", client_name).strip()
        target_dir = cgp_root / safe_name / subfolder

    target_dir.mkdir(parents=True, exist_ok=True)

    # Build output filename
    import datetime
    date_str = datetime.date.today().isoformat()
    type_slug = subfolder.replace("-", "").replace("/", "")
    if client_name:
        last_name = client_name.split()[-1]
        out_name = f"{date_str}_{type_slug}_{last_name}.docx"
    else:
        out_name = f"{date_str}_{type_slug}.docx"

    out_path = target_dir / out_name

    # Convert Markdown → docx
    try:
        _md_to_docx(md_path, out_path, config)
    except ImportError:
        sys.stderr.write("[cgp] python-docx manquant — relancer /charte ou /setup\n")
        return
    except Exception as exc:
        sys.stderr.write(f"[cgp] erreur conversion : {exc}\n")
        return

    # Delete source .md
    try:
        md_path.unlink()
    except Exception:
        pass

    sys.stderr.write(f"[cgp] → {out_path}\n")


if __name__ == "__main__":
    main()
